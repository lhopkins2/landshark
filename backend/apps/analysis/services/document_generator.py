import io
import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from fpdf import FPDF

_FONTS_DIR = Path(__file__).resolve().parent.parent.parent.parent / "fonts"
_FONT_REGULAR = str(_FONTS_DIR / "DejaVuSans.ttf")
_FONT_BOLD = str(_FONTS_DIR / "DejaVuSans-Bold.ttf")


def generate_docx(text: str, title: str = "") -> io.BytesIO:
    doc = DocxDocument()

    if title:
        doc.add_heading(title, level=1)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines):
                row_stripped = lines[i].strip()
                if _is_table_separator(row_stripped):
                    i += 1
                    continue
                if _is_table_row(row_stripped):
                    table_rows.append(_parse_table_row(row_stripped))
                    i += 1
                    continue
                break
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols, style="Table Grid")
                for col_idx, header in enumerate(table_rows[0]):
                    if _is_page_col(header):
                        table.columns[col_idx].width = Inches(0.55)
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_text in enumerate(row_data[:num_cols]):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = cell_text
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                    if row_idx == 0:
                        for col_idx in range(num_cols):
                            for run in table.cell(0, col_idx).paragraphs[0].runs:
                                run.bold = True
            continue
        else:
            doc.add_paragraph(stripped)

        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _clean_cell_text(text: str) -> str:
    """Collapse tabs, newlines, and runs of spaces into a single-line cell value."""
    text = re.sub(r"[\t\r\f\v]+", " ", text)
    text = re.sub(r"\n+", " ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [_clean_cell_text(c) for c in cells]


def _is_table_row(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[-:| ]+\|$", line))


def _is_page_col(header: str) -> bool:
    """True if the header looks like a page-number column (Doc Pg, Page, Pg, etc.)."""
    norm = re.sub(r"[^a-z]", "", header.lower())
    return norm in ("docpg", "page", "pages", "pg")


def strip_page_column(markdown_text: str) -> str:
    """Drop every "Doc Pg"-style column from any markdown table in the input.

    Identifies contiguous table blocks, removes cells whose header matches
    `_is_page_col`, and rebuilds the separator row to match the new column count.
    """
    out_lines: list[str] = []
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not _is_table_row(stripped):
            out_lines.append(line)
            i += 1
            continue

        block: list[list[str]] = []
        had_separator = False
        while i < len(lines):
            s = lines[i].strip()
            if _is_table_separator(s):
                had_separator = True
                i += 1
                continue
            if _is_table_row(s):
                block.append(_parse_table_row(s))
                i += 1
                continue
            break

        if not block:
            continue

        headers = block[0]
        drop_idx = {ci for ci, h in enumerate(headers) if _is_page_col(h)}
        kept = [
            [c for ci, c in enumerate(row) if ci not in drop_idx]
            for row in block
        ]

        out_lines.append("| " + " | ".join(kept[0]) + " |")
        if had_separator:
            out_lines.append("|" + "|".join(["---"] * len(kept[0])) + "|")
        for row in kept[1:]:
            out_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(out_lines)


def _compute_col_widths(headers: list[str], total_width: float) -> list[float]:
    """Compute column widths, giving page-number columns a narrow fixed width."""
    narrow = 18  # mm — just enough for "91-96"
    narrow_count = sum(1 for h in headers if _is_page_col(h))
    normal_count = len(headers) - narrow_count
    remaining = total_width - (narrow * narrow_count)
    normal_width = remaining / normal_count if normal_count else total_width / len(headers)
    return [narrow if _is_page_col(h) else normal_width for h in headers]


# Manual row layout. fpdf2's pdf.table() throws when a row's wrapped content can't
# fit on the remaining page, which forced an all-or-nothing fallback. Driving layout
# ourselves means we always know what fits and every row renders bordered correctly.
_TABLE_FONT_SIZE = 8
_LINE_HEIGHT_MM = 4.0
_CELL_PAD_X = 1.0
_CELL_PAD_Y = 1.0


def _measure_cell_height(pdf: FPDF, text: str, inner_w: float) -> float:
    """Wrapped height (mm) at the currently active font.

    Caller MUST set the table font first — fpdf2's dry_run uses whatever font is set.
    """
    if not text:
        return _LINE_HEIGHT_MM
    return pdf.multi_cell(
        w=inner_w,
        h=_LINE_HEIGHT_MM,
        text=text,
        dry_run=True,
        output="HEIGHT",
    )


def _measure_row_height(pdf: FPDF, row: list[str], col_widths: list[float]) -> float:
    heights = [
        _measure_cell_height(
            pdf,
            row[i] if i < len(row) else "",
            col_widths[i] - 2 * _CELL_PAD_X,
        )
        for i in range(len(col_widths))
    ]
    return max(heights) + 2 * _CELL_PAD_Y


def _draw_row(
    pdf: FPDF,
    row: list[str],
    col_widths: list[float],
    row_height: float,
    is_header: bool,
) -> None:
    """Draw a single row at pdf.y with uniform borders.

    Text goes through multi_cell (no border); a single rect() spanning row_height
    is then drawn over it so cells line up regardless of individual wrapped heights.
    """
    x0, y0 = pdf.l_margin, pdf.y
    pdf.set_font("DejaVu", "B" if is_header else "", _TABLE_FONT_SIZE)
    x = x0
    for i, w in enumerate(col_widths):
        text = row[i] if i < len(row) else ""
        pdf.set_xy(x + _CELL_PAD_X, y0 + _CELL_PAD_Y)
        pdf.multi_cell(
            w - 2 * _CELL_PAD_X,
            _LINE_HEIGHT_MM,
            text,
            border=0,
            align="L",
            new_x="RIGHT",
            new_y="TOP",
            max_line_height=_LINE_HEIGHT_MM,
        )
        pdf.rect(x, y0, w, row_height)
        x += w
    pdf.set_xy(x0, y0 + row_height)


def _wrap_text_to_height(pdf: FPDF, text: str, inner_w: float, budget_h: float) -> list[str]:
    """Greedy split of `text` into chunks each measuring <= budget_h at inner_w.

    Caller must have set the font used for measurement and drawing.
    """
    words = text.split()
    if not words:
        return [text]
    chunks: list[str] = []
    cur = ""
    for word in words:
        trial = (cur + " " + word).strip() if cur else word
        h = _measure_cell_height(pdf, trial, inner_w)
        if h > budget_h and cur:
            chunks.append(cur)
            cur = word
        else:
            cur = trial
    if cur:
        chunks.append(cur)
    return chunks or [text]


def _split_overflow_row(
    pdf: FPDF,
    row: list[str],
    col_widths: list[float],
    max_height: float,
) -> list[list[str]]:
    """Split the tallest cell across continuation rows so each fits in max_height.

    Other cells appear once on the first row; later rows show '(cont'd)' in
    column 0 and blanks elsewhere. Caller must have set the table font.
    """
    inner = [w - 2 * _CELL_PAD_X for w in col_widths]
    heights = [_measure_cell_height(pdf, row[i], inner[i]) for i in range(len(row))]
    worst = heights.index(max(heights))
    # Headroom: vertical padding + a line for "(cont'd)" on continuation rows.
    budget = max_height - 2 * _CELL_PAD_Y - _LINE_HEIGHT_MM
    pieces = _wrap_text_to_height(pdf, row[worst], inner[worst], budget)
    parts: list[list[str]] = []
    for idx, piece in enumerate(pieces):
        new_row = list(row)
        new_row[worst] = piece
        if idx > 0:
            for j in range(len(new_row)):
                new_row[j] = "(cont'd)" if j == 0 else ""
            new_row[worst] = piece
        parts.append(new_row)
    return parts


def _render_pdf_table(pdf: FPDF, rows: list[list[str]]) -> None:
    """Render a markdown-style table by laying out rows manually.

    Redraws the header at the top of each page. A row taller than an empty
    page's usable area is split across continuation rows on the tallest cell.
    """
    if not rows or len(rows) < 2:
        return

    headers, *data_rows = rows
    if not headers:
        return
    col_widths = _compute_col_widths(headers, pdf.epw)

    pdf.set_font("DejaVu", "B", _TABLE_FONT_SIZE)
    header_h = _measure_row_height(pdf, headers, col_widths)

    def _ensure_space(row_h: float) -> None:
        if pdf.y + row_h > pdf.h - pdf.b_margin:
            pdf.add_page()
            _draw_row(pdf, headers, col_widths, header_h, is_header=True)

    # Pagination is driven manually here; fpdf2's auto-break races with multi_cell mid-row.
    prev_auto_break = pdf.auto_page_break
    prev_b_margin = pdf.b_margin
    pdf.set_auto_page_break(auto=False, margin=prev_b_margin)
    try:
        _draw_row(pdf, headers, col_widths, header_h, is_header=True)

        usable_full_page = pdf.h - pdf.t_margin - prev_b_margin - header_h
        num_cols = len(col_widths)

        for raw in data_rows:
            row = list(raw[:num_cols]) + [""] * max(0, num_cols - len(raw))
            pdf.set_font("DejaVu", "", _TABLE_FONT_SIZE)
            row_h = _measure_row_height(pdf, row, col_widths)

            if row_h > usable_full_page:
                for part in _split_overflow_row(pdf, row, col_widths, usable_full_page):
                    pdf.set_font("DejaVu", "", _TABLE_FONT_SIZE)
                    part_h = _measure_row_height(pdf, part, col_widths)
                    _ensure_space(part_h)
                    _draw_row(pdf, part, col_widths, part_h, is_header=False)
            else:
                _ensure_space(row_h)
                _draw_row(pdf, row, col_widths, row_h, is_header=False)
    finally:
        pdf.set_auto_page_break(auto=prev_auto_break, margin=prev_b_margin)

    pdf.set_font("DejaVu", "", 10)


def generate_pdf(text: str, title: str = "") -> io.BytesIO:
    pdf = FPDF()
    pdf.add_font("DejaVu", "", _FONT_REGULAR)
    pdf.add_font("DejaVu", "B", _FONT_BOLD)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    if title:
        pdf.set_font("DejaVu", "B", 14)
        pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    pdf.set_font("DejaVu", "", 10)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            pdf.ln(4)
            i += 1
            continue

        if stripped.startswith("### "):
            pdf.set_font("DejaVu", "B", 11)
            pdf.cell(0, 7, stripped[4:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
        elif stripped.startswith("## "):
            pdf.set_font("DejaVu", "B", 12)
            pdf.cell(0, 8, stripped[3:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
        elif stripped.startswith("# "):
            pdf.set_font("DejaVu", "B", 14)
            pdf.cell(0, 10, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("DejaVu", "", 10)
        elif _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines):
                row_stripped = lines[i].strip()
                if _is_table_separator(row_stripped):
                    i += 1
                    continue
                if _is_table_row(row_stripped):
                    table_rows.append(_parse_table_row(row_stripped))
                    i += 1
                    continue
                break
            _render_pdf_table(pdf, table_rows)
            continue
        else:
            pdf.multi_cell(0, 5, stripped, new_x="LMARGIN", new_y="NEXT")

        i += 1

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def generate_document(text: str, output_format: str, title: str = "") -> io.BytesIO:
    if output_format == "docx":
        return generate_docx(text, title)
    return generate_pdf(text, title)


