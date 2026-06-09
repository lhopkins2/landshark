"""Prepare an uploaded DOCX so it can be filled by the COT template renderer.

Shops upload their *existing* COT form (plain Word doc with labels and an empty
instrument table — no template syntax). This module injects the docxtpl/Jinja2
placeholders automatically by recognizing:

  * the instrument table (a table row whose cells match known column headers), and
  * header-field labels (TAX ID, TRACT, RECORD OWNER, ADDRESS, ...).

The result is a docx the renderer can fill. If the doc already contains template
syntax (a power user or our starter template), it's left untouched.

Recognizing-not-requiring is the whole point: a shop should be able to upload the
form they already use and have it Just Work, without learning Jinja.
"""

import copy
import io
import re
from typing import Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Instrument-table columns: (match-substrings, placeholder text). The first
# column whose normalized header contains any substring wins.
_COLUMN_RULES: list[tuple[tuple[str, ...], str]] = [
    (("caption", "document", "instrument"), "{{ inst.caption }}"),
    (("bk", "book", "reception", "recording"), "{{ inst.book_page }}\n{{ inst.reception }}"),
    (("date recorded", "recorded", "date"), "{{ inst.recording_date }}"),
    (("grantor",), "{{ inst.grantor }}"),
    (("grantee",), "{{ inst.grantee }}"),
    (("legal", "comment"), "{{ inst.comments }}"),
]

# Header-field labels: (match-substring, placeholder). Checked most-specific first.
_HEADER_RULES: list[tuple[str, str]] = [
    ("tax id", "{{ tax_id }}"),
    ("parcel", "{{ tax_id }}"),
    ("tract", "{{ tract_number }}"),
    ("record owner", "{{ record_holder }}"),
    ("record holder", "{{ record_holder }}"),
    ("begin search", "{{ begin_search_date }}"),
    ("end search", "{{ end_search_date }}"),
    ("title agent", "{{ title_agent }}"),
    ("description", "{{ legal_description }}"),
    ("acres", "{{ acres }}"),
    ("address", "{{ property_address }}"),
]

# Substrings that indicate a row is the instrument-table header.
_INSTRUMENT_HEADER_HINTS = ("grantor", "grantee")


class TemplatePreparationError(Exception):
    """Raised when the uploaded docx has no recognizable instrument table."""


def _norm(text: str) -> str:
    """Lowercase, collapse whitespace, drop punctuation for fuzzy label matching."""
    return re.sub(r"[^a-z0-9 ]", "", (text or "").lower()).strip()


def _distinct_cells(row: Any) -> list[Any]:
    """Return one cell per merged region (python-docx repeats merged cells)."""
    seen: list[Any] = []
    out: list[Any] = []
    for cell in row.cells:
        if cell._tc not in seen:
            seen.append(cell._tc)
            out.append(cell)
    return out


def _set_cell_text(cell: Any, text: str) -> None:
    """Replace a cell's content with `text`, splitting on \\n into paragraphs.

    Keeps the FIRST paragraph's run formatting (font/size) by reusing it; clears
    any extra paragraphs. Cell-level formatting (borders, shading) is untouched.
    """
    lines = text.split("\n")
    # Reuse the first paragraph; clear its runs and write the first line.
    first_p = cell.paragraphs[0]
    # Preserve a run's formatting if one exists.
    template_run = first_p.runs[0] if first_p.runs else None
    for r in list(first_p.runs):
        r.text = ""
    if template_run is not None:
        template_run.text = lines[0]
    else:
        first_p.add_run(lines[0])
    # Remove any paragraphs beyond the first, then add one per remaining line.
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)
    for line in lines[1:]:
        p = cell.add_paragraph()
        run = p.add_run(line)
        if template_run is not None:
            run.font.size = template_run.font.size
            run.font.name = template_run.font.name


def _append_to_cell(cell: Any, text: str) -> None:
    """Append ` text` to a cell's existing first paragraph (keeps the label)."""
    p = cell.paragraphs[0]
    template_run = p.runs[0] if p.runs else None
    run = p.add_run(" " + text)
    if template_run is not None:
        run.font.size = template_run.font.size
        run.font.name = template_run.font.name
        run.bold = False


def _is_empty(cell: Any) -> bool:
    return not (cell.text or "").strip()


def _column_placeholder(header_text: str) -> str | None:
    norm = _norm(header_text)
    for substrings, placeholder in _COLUMN_RULES:
        if any(s in norm for s in substrings):
            return placeholder
    return None


def _header_placeholder(label_text: str) -> str | None:
    norm = _norm(label_text)
    for substring, placeholder in _HEADER_RULES:
        if substring in norm:
            return placeholder
    return None


def _looks_like_instrument_header(row: Any) -> bool:
    cells_norm = [_norm(c.text) for c in _distinct_cells(row)]
    joined = " ".join(cells_norm)
    return all(hint in joined for hint in _INSTRUMENT_HEADER_HINTS)


def _iter_all_tables(doc_or_cell: Any) -> Any:
    """Yield every table in the document, recursing into nested table cells."""
    tables = list(getattr(doc_or_cell, "tables", []))
    for tbl in tables:
        yield tbl
        for row in tbl.rows:
            for cell in row.cells:
                yield from _iter_all_tables(cell)


def _clone_marker_row(data_tr: Any, tag: str) -> Any:
    """Deep-copy a row's <w:tr>, blank all text, and write `tag` into its first cell."""
    new_tr = copy.deepcopy(data_tr)
    # Blank every text node in the clone.
    for t in new_tr.iter(qn("w:t")):
        t.text = ""
    # Write the tag into the first <w:t> we can find (create one if needed).
    first_t = next(iter(new_tr.iter(qn("w:t"))), None)
    if first_t is None:
        # Build a minimal run/paragraph in the first cell.
        first_tc = next(iter(new_tr.iter(qn("w:tc"))))
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = tag
        r.append(t)
        p.append(r)
        first_tc.append(p)
    else:
        first_t.text = tag
        first_t.set(qn("xml:space"), "preserve")
    return new_tr


def _inject_instrument_table(doc: Any) -> bool:
    """Find the instrument table, map its columns, and wrap the data row in a loop.

    Returns True if an instrument table was found and injected.
    """
    for tbl in _iter_all_tables(doc):
        for ri, row in enumerate(tbl.rows):
            if not _looks_like_instrument_header(row):
                continue
            header_cells = _distinct_cells(row)
            # The data row is the next row; if there isn't one, add a blank one.
            if ri + 1 >= len(tbl.rows):
                tbl.add_row()
            data_row = tbl.rows[ri + 1]
            data_cells = _distinct_cells(data_row)

            n = min(len(header_cells), len(data_cells))
            for ci in range(n):
                placeholder = _column_placeholder(header_cells[ci].text)
                if placeholder:
                    _set_cell_text(data_cells[ci], placeholder)

            # Wrap the data row in {%tr for %} / {%tr endfor %} marker rows.
            data_tr = data_row._tr
            for_tr = _clone_marker_row(data_tr, "{%tr for inst in instruments %}")
            endfor_tr = _clone_marker_row(data_tr, "{%tr endfor %}")
            data_tr.addprevious(for_tr)
            data_tr.addnext(endfor_tr)
            return True
    return False


def _inject_header_fields(doc: Any) -> int:
    """Inject value placeholders next to recognized header-field labels.

    Returns the count of fields injected. Best-effort: skips the instrument
    table rows and any cell that already has a placeholder.
    """
    injected = 0
    for tbl in _iter_all_tables(doc):
        header_row_idx = None
        for ri, row in enumerate(tbl.rows):
            if _looks_like_instrument_header(row):
                header_row_idx = ri
                break

        for ri, row in enumerate(tbl.rows):
            # Don't inject header-field values into the instrument table or below it.
            if header_row_idx is not None and ri >= header_row_idx:
                break
            cells = _distinct_cells(row)
            used: set[int] = set()
            for idx, cell in enumerate(cells):
                if idx in used:
                    continue
                placeholder = _header_placeholder(cell.text)
                if not placeholder or "{{" in cell.text or "{%" in cell.text:
                    continue
                # 1) Prefer an empty, non-label cell to the right in the same row.
                target = None
                for j in range(idx + 1, len(cells)):
                    if _is_empty(cells[j]) and _header_placeholder(cells[j].text) is None:
                        target = cells[j]
                        used.add(j)
                        break
                if target is not None:
                    _set_cell_text(target, placeholder)
                else:
                    # 2) Fall back to appending the placeholder after the label.
                    _append_to_cell(cell, placeholder)
                injected += 1
    return injected


def _append_notes_section(doc: Any) -> None:
    """Append a docxtpl-looped Notes section to the bottom of the document body.

    Mirrors the default renderer's "Notes" block (gaps, ambiguities, failed
    pages). Wrapped in `{%p if notes %}` so it renders nothing when the analysis
    produced no notes — no empty heading.
    """
    from docx.shared import Pt

    doc.add_paragraph("{%p if notes %}")
    heading = doc.add_paragraph()
    run = heading.add_run("Notes")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    doc.add_paragraph("{%p for note in notes %}")
    item = doc.add_paragraph()
    item_run = item.add_run("• {{ note.text }}{% if note.page %} (p.{{ note.page }}){% endif %}")
    item_run.font.size = Pt(10)
    item_run.font.name = "Arial"
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")


def prepare_uploaded_template(docx_bytes: bytes) -> bytes:
    """Return docx_bytes with placeholders injected, ready for the renderer.

    Raises TemplatePreparationError if no instrument table can be found.
    If the doc already contains template syntax, it's returned unchanged.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    # Already templated? Leave it alone.
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in _iter_all_tables(doc):
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    if "{{" in full_text or "{%" in full_text:
        return docx_bytes

    found_table = _inject_instrument_table(doc)
    if not found_table:
        raise TemplatePreparationError(
            "Couldn't find an instrument table in this template. Make sure it has a "
            "table whose header row includes columns like Grantor and Grantee."
        )
    _inject_header_fields(doc)
    _append_notes_section(doc)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
