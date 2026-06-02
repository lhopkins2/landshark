import io
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

# 100 DPI keeps peak memory and per-page token cost reasonable while staying
# readable for stamps, signatures, and cursive (~720×936px on US-Letter).
# 150 pages is a practical cap; beyond that the provider context window is the real limit.
IMAGE_DPI = 100
MAX_PAGES_REDUCED = 150
# JPEG quality 85 keeps text fully legible (artifacts only appear on photo-style
# regions) while shrinking page bytes ~5-10× vs PNG. PNG losslessly preserved scanner
# grain we didn't need, blowing past Anthropic's 32 MB request body cap on long deeds.
JPEG_QUALITY = 85


def extract_text_from_file(file_field: Any) -> str:
    """Extract text from a PDF, DOCX, or TXT file.

    `file_field` is duck-typed: anything with `.name`, `.open(mode)`, `.read()`,
    `.close()` (Django `FieldFile` in production).
    """
    filename = file_field.name.lower()

    if filename.endswith(".pdf"):
        return _extract_pdf_text(file_field)
    elif filename.endswith(".docx"):
        return _extract_docx_text(file_field)
    elif filename.endswith(".txt"):
        file_field.open("rb")
        content: str = file_field.read().decode("utf-8", errors="replace")
        file_field.close()
        return content
    else:
        raise ValueError(f"Unsupported file type: {Path(filename).suffix}")


def _extract_pdf_text(file_field: Any) -> str:
    import pymupdf

    file_field.open("rb")
    data = file_field.read()
    file_field.close()

    doc = pymupdf.open(stream=data, filetype="pdf")

    text_parts = []
    total_meaningful_chars = 0
    for i, page in enumerate(doc):
        text = page.get_text()
        text_parts.append(f"--- Page {i + 1} ---\n{text}")
        for line in text.split("\n"):
            stripped = line.strip()
            if len(stripped) <= 20:
                continue
            # Drop browser chrome (URLs, "page X of Y") so it doesn't dominate the count.
            if re.search(r"https?://|\.com|\.aspx|\.pdf|\d+\s+of\s+\d+", stripped, re.IGNORECASE):
                continue
            total_meaningful_chars += len(stripped)

    # Scanned PDFs are handled by the vision pipeline (render_pdf_pages), not OCR here.

    doc.close()
    return "\n".join(text_parts)


def render_pdf_pages(
    file_field: Any,
    dpi: int = IMAGE_DPI,
    max_pages: int | None = None,
    page_indexes: Iterable[int] | None = None,
) -> tuple[list[tuple[int, bytes]], int]:
    """Render PDF pages as JPEGs for vision analysis.

    Returns (list of (page_number, image_bytes), total_pages). page_number is 1-indexed.
    If page_indexes is given, only those pages are rendered (max_pages is ignored).
    """
    import pymupdf

    file_field.open("rb")
    data = file_field.read()
    file_field.close()

    doc = pymupdf.open(stream=data, filetype="pdf")
    total_pages = len(doc)

    if page_indexes is not None:
        seen = set()
        target = []
        for p in page_indexes:
            try:
                p = int(p)
            except (TypeError, ValueError):
                continue
            if 1 <= p <= total_pages and p not in seen:
                seen.add(p)
                target.append(p)
    else:
        render_count = min(total_pages, max_pages) if max_pages else total_pages
        target = list(range(1, render_count + 1))

    pages = []
    for page_num in target:
        page = doc[page_num - 1]
        pixmap = page.get_pixmap(dpi=dpi)
        image_bytes = pixmap.tobytes("jpeg", jpg_quality=JPEG_QUALITY)
        pages.append((page_num, image_bytes))

    doc.close()
    return pages, total_pages


def _extract_table_rows(table: Any) -> list[str]:
    """Flatten a docx table to text rows, recursing into nested tables."""
    rows: list[str] = []
    for row in table.rows:
        has_nested = any(cell.tables for cell in row.cells)
        if has_nested:
            for cell in row.cells:
                for nested_table in cell.tables:
                    rows.extend(_extract_table_rows(nested_table))
        else:
            # python-docx returns the same text in every merged cell — dedupe adjacent dupes.
            seen = []
            for cell in row.cells:
                text = cell.text.strip().replace("\n", " ")
                if not seen or text != seen[-1]:
                    seen.append(text)
            row_text = " | ".join(seen)
            if row_text.strip(" |"):
                rows.append(row_text)
    return rows


def _extract_docx_text(file_field: Any) -> str:
    from docx import Document as DocxDocument

    file_field.open("rb")
    data = file_field.read()
    file_field.close()

    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        parts.extend(_extract_table_rows(table))

    return "\n".join(parts)
